#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
EZONE MES 업무 프로세스 플로우 문서 생성기
Word(.docx) 형식으로 시각적 플로우차트 포함 문서 생성
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os, io

# ─── 색상 상수 ───
C_PRIMARY   = RGBColor(0x1B, 0x3A, 0x5C)   # 남색 (제목)
C_ACCENT    = RGBColor(0x2E, 0x86, 0xC1)   # 파랑 (강조)
C_SUCCESS   = RGBColor(0x27, 0xAE, 0x60)   # 녹색
C_WARNING   = RGBColor(0xF3, 0x9C, 0x12)   # 주황
C_DANGER    = RGBColor(0xE7, 0x4C, 0x3C)   # 빨강
C_GRAY      = RGBColor(0x7F, 0x8C, 0x8D)   # 회색
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK     = RGBColor(0x00, 0x00, 0x00)
C_LIGHT_BG  = RGBColor(0xEB, 0xF5, 0xFB)   # 연한 파랑 배경
C_STEP_BG   = "D6EAF8"
C_DECISION_BG = "FDEBD0"
C_ROLE_BG   = "E8F8F5"

def set_cell_bg(cell, color_hex):
    """셀 배경색 설정"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_margins(cell, top=50, bottom=50, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:start w:w="{left}" w:type="dxa"/>'
        f'  <w:end w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcMar_old = tcPr.find(qn('w:tcMar'))
    if tcMar_old is not None:
        tcPr.remove(tcMar_old)
    tcPr.append(tcMar)

def add_styled_paragraph(doc, text, style='Normal', font_size=None, bold=False,
                          color=None, alignment=None, space_before=None, space_after=None):
    p = doc.add_paragraph()
    if style and style != 'Normal':
        p.style = doc.styles[style]
    run = p.add_run(text)
    if font_size:
        run.font.size = Pt(font_size)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def make_flow_table(doc, steps, title=None):
    """
    시각적 플로우차트를 테이블로 구현
    steps: list of dict { 'no', 'name', 'actor', 'action', 'system', 'output', 'type' }
    type: 'process' | 'decision' | 'start' | 'end'
    """
    if title:
        add_styled_paragraph(doc, title, font_size=13, bold=True, color=C_PRIMARY,
                             space_before=12, space_after=6)

    # 헤더
    headers = ['순서', '업무단계', '담당자/역할', '실무 수행내용', '시스템 처리', '산출물/결과']
    col_widths = [Cm(1.2), Cm(2.8), Cm(2.5), Cm(5.5), Cm(3.5), Cm(3.0)]

    table = doc.add_table(rows=1 + len(steps), cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 헤더 행
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = C_WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, "1B3A5C")
        set_cell_margins(cell, 40, 40, 40, 40)

    # 데이터 행
    for row_idx, step in enumerate(steps):
        row = table.rows[row_idx + 1]

        # 순서 번호 (화살표 포함)
        c0 = row.cells[0]
        c0.text = ''
        p0 = c0.paragraphs[0]
        if step.get('type') == 'decision':
            symbol = f"  {step['no']}"
        else:
            symbol = f"{step['no']}"
        run0 = p0.add_run(symbol)
        run0.font.size = Pt(9)
        run0.bold = True
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 업무단계
        c1 = row.cells[1]
        c1.text = ''
        p1 = c1.paragraphs[0]
        # 플로우 아이콘
        if step.get('type') == 'decision':
            icon = '\u25C7 '  # ◇ 다이아몬드
            bg = C_DECISION_BG
        elif step.get('type') == 'start':
            icon = '\u25B6 '  # ▶
            bg = "D5F5E3"
        elif step.get('type') == 'end':
            icon = '\u25A0 '  # ■
            bg = "FADBD8"
        else:
            icon = '\u25A1 '  # □
            bg = C_STEP_BG

        r1 = p1.add_run(icon + step['name'])
        r1.font.size = Pt(9)
        r1.bold = True
        set_cell_bg(c1, bg)

        # 담당자
        c2 = row.cells[2]
        c2.text = ''
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(step.get('actor', ''))
        r2.font.size = Pt(8.5)
        set_cell_bg(c2, C_ROLE_BG)

        # 실무 수행내용
        c3 = row.cells[3]
        c3.text = ''
        p3 = c3.paragraphs[0]
        r3 = p3.add_run(step.get('action', ''))
        r3.font.size = Pt(8.5)

        # 시스템 처리
        c4 = row.cells[4]
        c4.text = ''
        p4 = c4.paragraphs[0]
        r4 = p4.add_run(step.get('system', ''))
        r4.font.size = Pt(8.5)
        r4.font.color.rgb = C_ACCENT

        # 산출물
        c5 = row.cells[5]
        c5.text = ''
        p5 = c5.paragraphs[0]
        r5 = p5.add_run(step.get('output', ''))
        r5.font.size = Pt(8.5)

        for ci in range(6):
            set_cell_margins(row.cells[ci], 30, 30, 40, 40)

    # 열 너비 설정
    for i, w in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = w

    return table


def make_arrow_flow(doc, items, color_hex="2E86C1"):
    """
    화살표로 연결된 수평 플로우 다이어그램
    items: list of str (단계명)
    """
    cols = len(items) * 2 - 1  # 단계 + 화살표
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, item in enumerate(items):
        cell_idx = i * 2
        cell = table.rows[0].cells[cell_idx]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(item)
        run.font.size = Pt(9)
        run.bold = True
        run.font.color.rgb = C_WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, color_hex)
        set_cell_margins(cell, 50, 50, 30, 30)

        # 화살표
        if i < len(items) - 1:
            arrow_cell = table.rows[0].cells[cell_idx + 1]
            arrow_cell.text = ''
            pa = arrow_cell.paragraphs[0]
            ra = pa.add_run('  \u2192  ')
            ra.font.size = Pt(14)
            ra.bold = True
            ra.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
            pa.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 테이블 테두리 제거
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:start w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:end w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'</w:tcBorders>'
            )
            old = tcPr.find(qn('w:tcBorders'))
            if old is not None:
                tcPr.remove(old)
            tcPr.append(tcBorders)

    return table


def make_vertical_flow(doc, items):
    """
    세로 플로우차트 (화살표 ↓ 로 연결)
    items: list of dict { 'label', 'detail', 'type', 'color' }
    """
    for i, item in enumerate(items):
        color = item.get('color', C_STEP_BG)
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        cell.text = ''
        cell.width = Cm(14)

        # 단계명
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        shape_icon = {
            'start': '\u25B6 ',
            'process': '\u25A1 ',
            'decision': '\u25C7 ',
            'end': '\u25A0 ',
            'sub': '    \u2514 ',
        }.get(item.get('type', 'process'), '\u25A1 ')

        run = p.add_run(shape_icon + item['label'])
        run.bold = True
        run.font.size = Pt(11)

        if item.get('detail'):
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(item['detail'])
            r2.font.size = Pt(8.5)
            r2.font.color.rgb = C_GRAY

        set_cell_bg(cell, color)
        set_cell_margins(cell, 60, 60, 80, 80)

        # 화살표 (마지막 아닌 경우)
        if i < len(items) - 1:
            arrow_p = doc.add_paragraph()
            arrow_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            arrow_p.paragraph_format.space_before = Pt(1)
            arrow_p.paragraph_format.space_after = Pt(1)
            ar = arrow_p.add_run('\u2193')
            ar.font.size = Pt(16)
            ar.bold = True
            ar.font.color.rgb = C_ACCENT


def make_info_box(doc, title, content_lines, bg_color="EBF5FB"):
    """정보 박스"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.text = ''
    cell.width = Cm(16)

    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = C_PRIMARY

    for line in content_lines:
        p2 = cell.add_paragraph()
        r2 = p2.add_run(line)
        r2.font.size = Pt(9)

    set_cell_bg(cell, bg_color)
    set_cell_margins(cell, 60, 60, 100, 100)


def make_role_matrix(doc, roles):
    """역할 매트릭스 테이블"""
    table = doc.add_table(rows=1 + len(roles), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    headers = ['역할', '담당업무', '주요 권한', '시스템 메뉴']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = C_WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, "1B3A5C")

    for idx, role in enumerate(roles):
        row = table.rows[idx + 1]
        for ci, val in enumerate(role):
            cell = row.cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8.5)
            set_cell_margins(cell, 30, 30, 40, 40)
            if ci == 0:
                r.bold = True
                set_cell_bg(cell, C_ROLE_BG)


# ═══════════════════════════════════════════════════════
# 메인 문서 생성
# ═══════════════════════════════════════════════════════

doc = Document()

# ─── 기본 스타일 설정 ───
style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_style.font.name = 'Malgun Gothic'
    h_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    h_style.font.color.rgb = C_PRIMARY

# 페이지 설정 (A4 세로)
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

# ═══════════════════════════════════════════════════════
# 표지
# ═══════════════════════════════════════════════════════

for _ in range(4):
    doc.add_paragraph()

# 회사 로고/제목
p_company = doc.add_paragraph()
p_company.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_company = p_company.add_run('EZONE')
r_company.font.size = Pt(36)
r_company.bold = True
r_company.font.color.rgb = C_PRIMARY

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub = p_sub.add_run('Manufacturing Execution System')
r_sub.font.size = Pt(14)
r_sub.font.color.rgb = C_GRAY

doc.add_paragraph()

# 구분선
p_line = doc.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_line = p_line.add_run('\u2500' * 60)
r_line.font.color.rgb = C_ACCENT
r_line.font.size = Pt(10)

doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_title = p_title.add_run('업무 프로세스 플로우 가이드')
r_title.font.size = Pt(26)
r_title.bold = True
r_title.font.color.rgb = C_PRIMARY

p_title2 = doc.add_paragraph()
p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_title2 = p_title2.add_run('수주 접수부터 인수검사 완료까지\n전체 업무 흐름 및 시스템 운용 절차서')
r_title2.font.size = Pt(13)
r_title2.font.color.rgb = C_GRAY

for _ in range(6):
    doc.add_paragraph()

# 문서 정보 테이블
info_table = doc.add_table(rows=5, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ('문서번호', 'GDL_20260330_업무프로세스플로우가이드'),
    ('작 성 일', '2026년 03월 30일'),
    ('작 성 자', 'June (품질관리)'),
    ('적용범위', '방화구획 관통부 제품 (소켓, 플래싱, 틈새시트)'),
    ('관련시스템', 'EZONE MES v1.0'),
]
for i, (k, v) in enumerate(info_data):
    c0 = info_table.rows[i].cells[0]
    c0.text = ''
    p0 = c0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    r0.font.size = Pt(10)
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(c0, "EBF5FB")
    c0.width = Cm(3)

    c1 = info_table.rows[i].cells[1]
    c1.text = ''
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(v)
    r1.font.size = Pt(10)
    c1.width = Cm(10)

    for c in [c0, c1]:
        set_cell_margins(c, 40, 40, 60, 60)

# 결재란
doc.add_paragraph()
approval_table = doc.add_table(rows=3, cols=3)
approval_table.alignment = WD_TABLE_ALIGNMENT.CENTER
approval_table.style = 'Table Grid'
headers_a = ['작 성', '검 토', '승 인']
for i, h in enumerate(headers_a):
    cell = approval_table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(cell, "D6EAF8")
    cell.width = Cm(4)

for ri in range(1, 3):
    for ci in range(3):
        cell = approval_table.rows[ri].cells[ci]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if ri == 1:
            cell.height = Cm(2)
        else:
            r = p.add_run('     /     /     ')
            r.font.size = Pt(9)
            r.font.color.rgb = C_GRAY

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 목차
# ═══════════════════════════════════════════════════════

add_styled_paragraph(doc, '목  차', font_size=18, bold=True, color=C_PRIMARY,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

toc_items = [
    ('1.', '문서 개요', '3'),
    ('2.', '전체 업무 플로우 개관', '4'),
    ('3.', '역할 및 권한 매트릭스', '5'),
    ('4.', '[STEP 1] 수주 접수 및 등록', '6'),
    ('5.', '[STEP 2] BOM 전개 (자재소요량 산출)', '8'),
    ('6.', '[STEP 3] 발주서 생성 및 승인', '10'),
    ('7.', '[STEP 4] 자재 입고 등록', '12'),
    ('8.', '[STEP 5] 인수검사 실행', '14'),
    ('9.', '[STEP 6] 작업지시 및 생산', '17'),
    ('10.', '실제 운영 예시 (SO-260303-001)', '19'),
    ('11.', '부적합 처리 플로우', '22'),
    ('12.', '부록: 양식코드 매핑표', '23'),
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
    r1 = p.add_run(f'{num}  {title}')
    r1.font.size = Pt(11)
    if not num.startswith((' ', '  ')):
        r1.bold = True
    r2 = p.add_run(f'\t{page}')
    r2.font.size = Pt(11)
    r2.font.color.rgb = C_GRAY

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 1. 문서 개요
# ═══════════════════════════════════════════════════════

doc.add_heading('1. 문서 개요', level=1)

doc.add_heading('1.1 목적', level=2)
add_styled_paragraph(doc,
    '본 문서는 EZONE MES 시스템의 전체 업무 프로세스를 실무자 관점에서 설명한다. '
    '수주 접수부터 BOM 전개, 발주서 생성/승인, 자재 입고, 인수검사, 작업지시에 이르는 '
    '전 과정의 업무 흐름, 담당 역할, 시스템 입력 절차를 시각적으로 기술하여 '
    '신규 투입 인원 교육 및 품질인정 심사 대비 자료로 활용한다.',
    font_size=10, space_after=6)

doc.add_heading('1.2 적용범위', level=2)
scope_items = [
    '방화구획 관통부 제품: 방화소켓(벽체/입상), 방화플래싱(I형/Z형), 틈새복합시트',
    '대상 자재: 원자재(RM) 4종, 부자재(SM) 9종, 반제품(SA) 6종, 완제품(FP) 7종',
    '관련 사규: C-701 Rev.5, C-302 Rev.8, C-601 Rev.4, D-121~D-126',
]
for item in scope_items:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)

doc.add_heading('1.3 용어 정의', level=2)

terms = [
    ('BOM', 'Bill of Materials. 제품 구성 자재 목록 및 소요량'),
    ('LOT', '동일 조건 생산/입고된 자재의 추적 단위'),
    ('인수검사', '입고 자재의 품질 적합성을 검증하는 검사 (n=3, c=0)'),
    ('n=3, c=0', '3개 시료 측정, 불합격 허용 0개'),
    ('MES', 'Manufacturing Execution System. 생산실행시스템'),
    ('PR', 'Purchase Request. 구매발주요청서'),
]

term_table = doc.add_table(rows=1 + len(terms), cols=2)
term_table.style = 'Table Grid'
term_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['용어', '정의']):
    c = term_table.rows[0].cells[i]
    c.text = ''
    p = c.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = C_WHITE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(c, "1B3A5C")
for idx, (term, defn) in enumerate(terms):
    r0 = term_table.rows[idx+1].cells[0]
    r0.text = ''
    rr = r0.paragraphs[0].add_run(term)
    rr.bold = True
    rr.font.size = Pt(9)
    set_cell_bg(r0, "EBF5FB")
    r0.width = Cm(3)

    r1 = term_table.rows[idx+1].cells[1]
    r1.text = ''
    rr1 = r1.paragraphs[0].add_run(defn)
    rr1.font.size = Pt(9)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 2. 전체 업무 플로우 개관
# ═══════════════════════════════════════════════════════

doc.add_heading('2. 전체 업무 플로우 개관', level=1)

add_styled_paragraph(doc,
    '아래는 수주 접수부터 생산 완료까지의 전체 업무 플로우를 요약한 다이어그램이다. '
    '각 단계는 실무 입력 → 시스템 자동 처리 → 관리자 확인/승인의 흐름으로 구성된다.',
    font_size=10, space_after=10)

# 메인 플로우 다이어그램 (세로)
make_vertical_flow(doc, [
    {'label': 'STEP 1. 수주 접수/등록', 'detail': '영업팀 \u2192 수주정보 입력 (고객, 현장, 구조, 수량)', 'type': 'start', 'color': 'D5F5E3'},
    {'label': 'STEP 2. BOM 전개', 'detail': '시스템 자동 \u2192 치수 기반 26개 품목 소요량 산출', 'type': 'process', 'color': 'D6EAF8'},
    {'label': 'STEP 3. 발주서 생성/승인', 'detail': '자재팀 작성 \u2192 관리자 승인 \u2192 발주 확정', 'type': 'decision', 'color': 'FDEBD0'},
    {'label': 'STEP 4. 자재 입고', 'detail': '창고팀 \u2192 입고등록, LOT 자동부여, 인수검사 연결', 'type': 'process', 'color': 'D6EAF8'},
    {'label': 'STEP 5. 인수검사', 'detail': '품질팀 \u2192 양식별 측정값 입력, 자동 합격/불합격 판정', 'type': 'decision', 'color': 'FDEBD0'},
    {'label': 'STEP 6. 작업지시/생산', 'detail': '생산팀 \u2192 배합\u2192압출\u2192재단\u2192조립 CASCADE 작업지시', 'type': 'process', 'color': 'D6EAF8'},
    {'label': '출하/납품', 'detail': '완제품 출하 \u2192 고객 인도', 'type': 'end', 'color': 'FADBD8'},
])

doc.add_paragraph()

make_info_box(doc, '\u26A0  핵심 포인트: 자동화 연결', [
    '\u2022 BOM 전개 \u2192 발주서 자동 생성: 부족 자재(shortage_qty > 0)만 자동 필터링하여 PR 생성',
    '\u2022 입고 등록 \u2192 인수검사 자동 생성: LOT 부여와 동시에 품목별 양식 매칭, 검사항목 프리셋 자동 입력',
    '\u2022 인수검사 \u2192 자동 판정: 측정값 입력 시 기준치(cert_standard) 대비 자동 PASS/FAIL 판정',
    '\u2022 작업지시 \u2192 공정별 CASCADE: MIX(배합) \u2192 EXT(압출) \u2192 CUT(재단) \u2192 ASM(조립) 자동 생성',
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 3. 역할 및 권한 매트릭스
# ═══════════════════════════════════════════════════════

doc.add_heading('3. 역할 및 권한 매트릭스', level=1)

add_styled_paragraph(doc,
    '각 업무 단계별 담당 역할과 시스템 권한을 정의한다.',
    font_size=10, space_after=8)

make_role_matrix(doc, [
    ['영업팀 (실무)', '수주 접수, 고객/현장 정보 입력,\n구조 및 수량 등록', '수주 생성/수정', '수주관리 메뉴'],
    ['자재팀 (실무)', 'BOM 전개 실행, 발주서 작성,\n입고 등록', 'BOM 전개, PR 생성,\n입고 처리', '수주상세, 발주관리,\n입고등록'],
    ['품질팀 (실무)', '인수검사 측정값 입력,\n성적서 확인', '검사 측정값 입력,\n판정 확인', '인수검사 메뉴'],
    ['생산팀 (실무)', '작업지시 확인, 생산 실적 입력', '작업지시 조회,\n실적 등록', '작업지시/실적 메뉴'],
    ['관리자 (팀장)', '발주서 승인, 검사 결과 최종확인,\n부적합 처리 결정', '승인/반려 권한,\n상태 변경', '전체 메뉴 + 승인 권한'],
    ['시스템 (자동)', 'BOM 소요량 계산, LOT 자동부여,\n검사 자동판정, PR 자동생성', '자동 처리\n(사용자 개입 없음)', '백그라운드 처리'],
])

doc.add_paragraph()

# RACI 차트
add_styled_paragraph(doc, 'RACI 매트릭스 (R=실행, A=승인, C=협의, I=통보)', font_size=11, bold=True,
                     color=C_PRIMARY, space_before=10, space_after=6)

raci = doc.add_table(rows=8, cols=6)
raci.style = 'Table Grid'
raci.alignment = WD_TABLE_ALIGNMENT.CENTER
raci_headers = ['업무단계', '영업팀', '자재팀', '품질팀', '생산팀', '관리자']
raci_data = [
    ['수주 접수', 'R', 'I', 'I', 'I', 'A'],
    ['BOM 전개', 'I', 'R', 'C', 'I', 'I'],
    ['발주서 생성', 'I', 'R', 'I', 'I', 'A'],
    ['자재 입고', 'I', 'R', 'I', 'I', 'I'],
    ['인수검사', 'I', 'I', 'R', 'I', 'A'],
    ['작업지시', 'I', 'I', 'I', 'R', 'A'],
    ['생산 실적', 'I', 'I', 'C', 'R', 'A'],
]

for ci, h in enumerate(raci_headers):
    c = raci.rows[0].cells[ci]
    c.text = ''
    r = c.paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = C_WHITE
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(c, "1B3A5C")

for ri, row_data in enumerate(raci_data):
    for ci, val in enumerate(row_data):
        c = raci.rows[ri+1].cells[ci]
        c.text = ''
        r = c.paragraphs[0].add_run(val)
        r.font.size = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if val == 'R':
            r.bold = True
            r.font.color.rgb = C_ACCENT
        elif val == 'A':
            r.bold = True
            r.font.color.rgb = C_DANGER

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 4. STEP 1 - 수주 접수
# ═══════════════════════════════════════════════════════

doc.add_heading('4. [STEP 1] 수주 접수 및 등록', level=1)

add_styled_paragraph(doc, '4.1 업무 개요', font_size=12, bold=True, color=C_PRIMARY, space_before=6)
add_styled_paragraph(doc,
    '고객으로부터 발주서(엑셀 또는 서면)를 접수하면, 영업팀 실무자가 MES 시스템에 수주 정보를 등록한다. '
    '수주 등록 시 고객명, 현장명, 구조 종류, 관통부 치수, 수량을 입력하며, '
    '이 정보가 이후 BOM 전개의 기초 데이터가 된다.',
    font_size=10, space_after=6)

make_flow_table(doc, [
    {'no': '1-1', 'name': '발주서 접수', 'actor': '영업팀 실무자',
     'action': '고객 발주서(엑셀/서면) 접수\n- 구조 종류, 관통부 치수, 수량 확인\n- 납기일, 특이사항 확인',
     'system': '-', 'output': '고객 발주서 원본', 'type': 'start'},
    {'no': '1-2', 'name': '수주 기본정보 입력', 'actor': '영업팀 실무자',
     'action': 'MES [수주관리] 메뉴 접속\n- 고객명, 현장명, 납기일 입력\n- 발주서 파일 첨부 (선택)',
     'system': '수주번호 자동 생성\nSO-YYMMDD-NNN',
     'output': 'SO-260303-001\n(수주번호)', 'type': 'process'},
    {'no': '1-3', 'name': '구조/수량 상세 입력', 'actor': '영업팀 실무자',
     'action': '수주 상세 품목 입력:\n- 인정구조 선택 (예: VA-064)\n- 관통부 치수 (W x H) 입력\n- 세트 수량 입력\n- 구조별 반복 추가',
     'system': '인정구조 마스터 연동\n(certification_master)\n구조코드 자동 매칭',
     'output': '수주상세 품목\n(10개 라인)', 'type': 'process'},
    {'no': '1-4', 'name': '수주 확정', 'actor': '관리자 (팀장)',
     'action': '수주 내역 검토\n- 구조/치수/수량 정확성 확인\n- 납기일 실현 가능성 판단',
     'system': '상태: REGISTERED\n\u2192 CONFIRMED',
     'output': '수주 확정 통보', 'type': 'decision'},
], title='수주 접수 상세 플로우')

doc.add_paragraph()

make_info_box(doc, '\U0001F4CB  수주 입력 시 주의사항', [
    '\u2022 관통부 치수(W x H)는 mm 단위로 정확히 입력 (BOM 소요량 계산의 핵심 변수)',
    '\u2022 동일 구조라도 치수가 다르면 별도 라인으로 입력 (소켓 크기가 달라짐)',
    '\u2022 소켓 수량(install_qty)은 1개 또는 2개 (2개 시 중앙 브라켓 추가 산출)',
    '\u2022 입상용(HTG/HAG계열)과 벽체용(VA/VT/VS계열) 구분 필수 (높이 기준 상이)',
])

doc.add_page_break()

# 입력 화면 설명
add_styled_paragraph(doc, '4.2 수주 등록 입력 항목', font_size=12, bold=True, color=C_PRIMARY, space_before=6)

input_fields = doc.add_table(rows=11, cols=4)
input_fields.style = 'Table Grid'
input_fields.alignment = WD_TABLE_ALIGNMENT.CENTER

field_headers = ['입력 필드', '데이터 타입', '필수여부', '설명']
for i, h in enumerate(field_headers):
    c = input_fields.rows[0].cells[i]
    c.text = ''
    r = c.paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = C_WHITE
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(c, "1B3A5C")

field_data = [
    ('고객명', '텍스트', '필수', '고객 회사명 (예: ㈜하나로엔지니어링)'),
    ('현장명', '텍스트', '필수', '시공 현장명 (예: 인천검단101역세권C1현장)'),
    ('납기일', '날짜', '필수', 'YYYY-MM-DD 형식'),
    ('인정구조', '선택(드롭다운)', '필수', 'VA-064, VT-01, HTG-064 등 13종 중 선택'),
    ('관통부 W', '숫자(mm)', '필수', '관통부 가로 치수'),
    ('관통부 H', '숫자(mm)', '필수', '관통부 세로 치수'),
    ('소켓수량(N)', '숫자', '필수', '1 또는 2 (대형 관통부 시 2개)'),
    ('세트 수량', '숫자', '필수', '해당 구조의 주문 세트 수'),
    ('비고', '텍스트', '선택', '특이사항 기재'),
    ('첨부파일', '파일', '선택', '고객 발주서 원본 업로드'),
]

for idx, (f1, f2, f3, f4) in enumerate(field_data):
    row = input_fields.rows[idx+1]
    vals = [f1, f2, f3, f4]
    for ci, v in enumerate(vals):
        c = row.cells[ci]
        c.text = ''
        r = c.paragraphs[0].add_run(v)
        r.font.size = Pt(8.5)
        if ci == 2 and v == '필수':
            r.bold = True
            r.font.color.rgb = C_DANGER

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 5. STEP 2 - BOM 전개
# ═══════════════════════════════════════════════════════

doc.add_heading('5. [STEP 2] BOM 전개 (자재소요량 산출)', level=1)

add_styled_paragraph(doc, '5.1 업무 개요', font_size=12, bold=True, color=C_PRIMARY, space_before=6)
add_styled_paragraph(doc,
    '수주 확정 후, 자재팀 실무자가 [BOM 전개] 버튼을 클릭하면 시스템이 '
    '관통부 치수(W x H)를 기반으로 26개 품목의 소요량을 자동 계산한다. '
    '계산 결과는 현재 재고와 대비하여 부족 수량(shortage_qty)을 산출하며, '
    '이것이 발주서 생성의 기초가 된다.',
    font_size=10, space_after=6)

make_flow_table(doc, [
    {'no': '2-1', 'name': 'BOM 전개 실행', 'actor': '자재팀 실무자',
     'action': '수주 상세 화면에서 [BOM 전개] 클릭\n- 수주 품목별 치수 파라미터 확인',
     'system': 'POST /orders/:id/explode-bom\n치수 기반 calculateStructureBom() 호출',
     'output': '26개 품목 소요량', 'type': 'process'},
    {'no': '2-2', 'name': '소요량 자동 산출', 'actor': '시스템 (자동)',
     'action': '관통부 둘레 \u2192 소켓/플래싱/차열재 수량\n배합 원자재 역전개 (reverse explode)\n- SA-CUT \u2192 SA-EXT \u2192 SA-MIX \u2192 RM',
     'system': 'calculatePerimeter()\naddFlashingComponents()\nreverseExplode()',
     'output': 'BOM 트리 구조\n(4단계 계층)', 'type': 'process'},
    {'no': '2-3', 'name': '재고 대비 부족량 산출', 'actor': '시스템 (자동)',
     'action': '품목별 현재 재고(current_stock) 조회\n소요량 - 재고 = 부족량(shortage_qty)\n부족량 > 0인 품목이 발주 대상',
     'system': 'inventory_transaction 집계\norder_bom_result 저장',
     'output': '부족 자재 목록\n(13개 품목)', 'type': 'process'},
    {'no': '2-4', 'name': 'BOM 결과 확인', 'actor': '자재팀 실무자',
     'action': 'BOM 전개 결과 화면에서 확인:\n- 4개 카테고리(RM/SM/SA/FP) 품목별 소요량\n- 재고 대비 부족량\n- 산출 근거(calc_note) 확인',
     'system': '상태: CONFIRMED\n\u2192 BOM_EXPLODED',
     'output': 'BOM 전개 완료\n발주 준비', 'type': 'decision'},
], title='BOM 전개 상세 플로우')

doc.add_paragraph()

# BOM 계층 구조
add_styled_paragraph(doc, '5.2 BOM 계층 구조 (4단계)', font_size=12, bold=True, color=C_PRIMARY, space_before=8)

bom_tree = doc.add_table(rows=6, cols=5)
bom_tree.style = 'Table Grid'
bom_tree.alignment = WD_TABLE_ALIGNMENT.CENTER

bom_headers = ['단계', '카테고리', '코드 예시', '품목 예시', '산출 방식']
for i, h in enumerate(bom_headers):
    c = bom_tree.rows[0].cells[i]
    c.text = ''
    r = c.paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = C_WHITE
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(c, "1B3A5C")

bom_data = [
    ('1단계\n(완제품)', 'FP', 'FP-VA064\nFP-FL-I', '방화소켓(VA-064)\n플래싱(I형)', '구조별 1:1'),
    ('2단계\n(부자재)', 'SM', 'SM-GW-24\nSM-GI-I\nSM-BRK-TB', '글라스울(24K)\n아연도금강판(I형)\n소켓 브라켓(상/하)', '치수 기반 계산\n플래싱=강판+차열시트\n소켓=본체+브라켓'),
    ('3단계\n(반제품)', 'SA', 'SA-CUT-SK\nSA-EXT-5190\nSA-MIX-MB', '재단(소켓용)\n압출(5T-190)\n인정배합', '역전개(reverse explode)\nCUT\u2192EXT\u2192MIX'),
    ('4단계\n(원자재)', 'RM', 'RM-MB\nRM-EA\nRM-EG50', '난연컴파운드\nEVA-EA33045\n팽창흑연 #50', '배합비율 기반\n300kg/배치 기준'),
    ('구매대상', 'RM+SM', '13품목', '원자재 4종 + 부자재 9종', 'shortage_qty > 0\n자재만 발주'),
]

for ri, row_data in enumerate(bom_data):
    for ci, val in enumerate(row_data):
        c = bom_tree.rows[ri+1].cells[ci]
        c.text = ''
        r = c.paragraphs[0].add_run(val)
        r.font.size = Pt(8.5)
        if ci == 0:
            r.bold = True
            set_cell_bg(c, "D6EAF8")

doc.add_page_break()

# BOM 핵심 계산식
add_styled_paragraph(doc, '5.3 BOM 핵심 계산 로직', font_size=12, bold=True, color=C_PRIMARY, space_before=6)

make_info_box(doc, '\U0001F4D0  플래싱 수량 계산식 (calcFlashingWall)', [
    '1면 둘레 = (W + 250) x 2 + H x 2     (W: 관통부가로, H: 관통부세로)',
    '양면(내/외부) = 1면 둘레 x 2',
    '필요 장수 = 양면 둘레 / 1000mm(1장 길이) x 로스율(10%)',
    '최종 수량 = CEIL(필요 장수)   (올림 처리)',
    '',
    '예시: W=700, H=600 \u2192 1면=(700+250)x2+600x2=3100 \u2192 양면=6200',
    '\u2192 필요=6200/1000x1.1=6.82 \u2192 CEIL=7장/세트 x 7세트 = 49장',
], bg_color="FEF9E7")

doc.add_paragraph()

make_info_box(doc, '\U0001F4D0  소켓 브라켓 수량 계산식', [
    '브라켓(상/하) = 소켓 N개 x 4EA/소켓    (상하 각 2개씩)',
    '브라켓(중앙) = 소켓 N개 x 2EA/소켓      (대형 N>=2인 경우만)',
    '',
    '예시: VA-064 (N=1, 11세트) \u2192 상/하 = 1x4x11 = 44EA, 중앙 = 없음(N<2)',
    '예시: HTG-064 (N=2, 5세트) \u2192 상/하 = 2x4x5 = 40EA, 중앙 = 2x2x5 = 20EA',
], bg_color="FEF9E7")

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 6. STEP 3 - 발주서 생성 및 승인
# ═══════════════════════════════════════════════════════

doc.add_heading('6. [STEP 3] 발주서 생성 및 승인', level=1)

add_styled_paragraph(doc, '6.1 업무 개요', font_size=12, bold=True, color=C_PRIMARY, space_before=6)
add_styled_paragraph(doc,
    'BOM 전개 후 부족 자재에 대해 발주서(PR)를 생성한다. '
    '시스템이 부족량을 자동으로 발주서 품목에 반영하며, '
    '관리자 승인을 거쳐 공급업체에 발주를 확정한다.',
    font_size=10, space_after=6)

make_flow_table(doc, [
    {'no': '3-1', 'name': '발주서 자동 생성', 'actor': '자재팀 실무자',
     'action': '수주 상세 \u2192 [발주서 생성] 클릭\n- 부족 자재(RM+SM) 자동 추출\n- 공급업체명 입력',
     'system': 'POST /orders/:id/create-pr\nshortage_qty > 0 필터\nPR번호 자동생성',
     'output': 'PR-260330-002\n(13품목, DRAFT)', 'type': 'process'},
    {'no': '3-2', 'name': '발주 내역 검토', 'actor': '자재팀 실무자',
     'action': '자동 생성된 발주 품목 검토:\n- 품목코드, 품명, 수량, 단위 확인\n- 단가 입력 (선택)\n- 납기일 설정 (선택)',
     'system': '발주서 상세 화면',
     'output': '발주서 내역 확정', 'type': 'process'},
    {'no': '3-3', 'name': '발주서 제출', 'actor': '자재팀 실무자',
     'action': '[제출] 버튼 클릭\n\u2192 관리자 승인 요청',
     'system': 'PATCH /purchase-requests/:id/status\nDRAFT \u2192 SUBMITTED',
     'output': '승인 요청 통보', 'type': 'process'},
    {'no': '3-4', 'name': '발주서 승인', 'actor': '관리자 (팀장)',
     'action': '발주 내역 검토 후 승인/반려:\n- 수량 적정성 확인\n- 납기 가능 여부 확인\n- 예산 확인',
     'system': 'SUBMITTED \u2192 APPROVED\n또는 \u2192 DRAFT(반려)',
     'output': '승인 완료', 'type': 'decision'},
    {'no': '3-5', 'name': '발주 확정', 'actor': '자재팀 실무자',
     'action': '승인된 발주서를 공급업체에 전달\n(이메일/팩스/전화)',
     'system': 'APPROVED \u2192 ORDERED',
     'output': '발주 완료\n입고 대기', 'type': 'end'},
], title='발주서 생성/승인 상세 플로우')

doc.add_paragraph()

# 상태 전이도
add_styled_paragraph(doc, '6.2 발주서 상태 전이도', font_size=12, bold=True, color=C_PRIMARY, space_before=8)

make_arrow_flow(doc, ['DRAFT\n(초안)', 'SUBMITTED\n(제출)', 'APPROVED\n(승인)', 'ORDERED\n(발주)', 'RECEIVED\n(입고완료)'])

doc.add_paragraph()

p_note = doc.add_paragraph()
r_note = p_note.add_run('* CANCELLED(취소): 어느 상태에서든 취소 가능. 취소 후 재생성 필요.')
r_note.font.size = Pt(9)
r_note.font.color.rgb = C_GRAY
r_note.italic = True

doc.add_paragraph()

make_info_box(doc, '\U0001F6AB  중복 방지 규칙', [
    '\u2022 동일 수주에 대해 CANCELLED가 아닌 발주서가 이미 존재하면 신규 생성 불가',
    '\u2022 기존 발주서 취소(CANCELLED) 후 재생성해야 함',
    '\u2022 발주서 취소 시 이미 입고된 품목의 LOT/검사 기록은 유지됨',
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 7. STEP 4 - 자재 입고
# ═══════════════════════════════════════════════════════

doc.add_heading('7. [STEP 4] 자재 입고 등록', level=1)

add_styled_paragraph(doc, '7.1 업무 개요', font_size=12, bold=True, color=C_PRIMARY, space_before=6)
add_styled_paragraph(doc,
    '공급업체로부터 자재가 입고되면, 창고팀(또는 자재팀)이 입고 등록을 수행한다. '
    '입고 등록 시 시스템이 자동으로 자체 LOT번호를 부여하고, '
    '해당 품목에 맞는 인수검사 양식을 자동 생성하여 검사 대기 상태로 연결한다.',
    font_size=10, space_after=6)

make_flow_table(doc, [
    {'no': '4-1', 'name': '자재 도착 확인', 'actor': '창고팀 실무자',
     'action': '공급업체 납품서/거래명세서 수령\n- 품목, 수량, LOT번호 확인\n- 외관 이상 여부 1차 확인',
     'system': '-',
     'output': '납품서 원본', 'type': 'start'},
    {'no': '4-2', 'name': '입고 정보 입력', 'actor': '창고팀 실무자',
     'action': 'MES [입고등록] 메뉴 접속\n- 발주서 선택 \u2192 품목 선택\n- 입고수량 입력\n- 공급처 LOT번호 입력\n- 검사자 지정',
     'system': 'POST /purchase-requests/:prId\n/items/:priId/receive',
     'output': '입고 데이터', 'type': 'process'},
    {'no': '4-3', 'name': 'LOT 자동 부여', 'actor': '시스템 (자동)',
     'action': '시스템 자동 처리:\n\u2460 입고LOT: IN-YYMMDD-NNN 생성\n\u2461 자체LOT: EZ-YYMMDD-NNN 생성\n\u2462 공급처LOT 연결',
     'system': 'lot_transaction INSERT\nlot_type = "IN"\nstatus = "ACTIVE"',
     'output': 'LOT번호 부여\n(추적성 확보)', 'type': 'process'},
    {'no': '4-4', 'name': '인수검사 자동 생성', 'actor': '시스템 (자동)',
     'action': '품목코드 \u2192 양식코드 자동 매칭:\n- SM-GW-24 \u2192 D122-1(그라스울 8항목)\n- SM-CW-96 \u2192 D124-1(세라믹울 7항목)\n- SM-GI-I \u2192 D121-2(강재류 9항목)\n- SM-BRK-TB \u2192 D121-2(강재류 9항목)',
     'system': 'inspection INSERT\ninspection_detail INSERT\nstatus = "PENDING"\nINS-NNNN 자동번호',
     'output': '인수검사 자동 생성\n(검사 대기)', 'type': 'process'},
    {'no': '4-5', 'name': '입고 상태 업데이트', 'actor': '시스템 (자동)',
     'action': '입고수량 vs 발주수량 비교:\n- 전량 입고: RECEIVED\n- 일부 입고: PARTIAL\n전체 입고 완료 시 PR 상태도 갱신',
     'system': 'receiving_status 갱신\nPR status 연동',
     'output': '입고 완료\n\u2192 검사 대기', 'type': 'end'},
], title='자재 입고 상세 플로우')

doc.add_paragraph()

# LOT 체계
add_styled_paragraph(doc, '7.2 LOT번호 체계 (C-302 Rev.8 연동)', font_size=12, bold=True, color=C_PRIMARY, space_before=8)

lot_table = doc.add_table(rows=4, cols=4)
lot_table.style = 'Table Grid'
lot_table.alignment = WD_TABLE_ALIGNMENT.CENTER

lot_headers = ['LOT 종류', '형식', '예시', '용도']
for i, h in enumerate(lot_headers):
    c = lot_table.rows[0].cells[i]
    c.text = ''
    r = c.paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = C_WHITE
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(c, "1B3A5C")

lot_data = [
    ('입고 LOT', 'IN-YYMMDD-NNN', 'IN-260330-005', '입고 이력 추적'),
    ('자체 LOT', 'EZ-YYMMDD-NNN', 'EZ-260330-005', '인수검사 LOT (C-302 연동)'),
    ('공급처 LOT', '업체 자체 번호', 'GW24-2603-01', '공급업체 추적 (역추적)'),
]
for ri, (t1, t2, t3, t4) in enumerate(lot_data):
    for ci, val in enumerate([t1, t2, t3, t4]):
        c = lot_table.rows[ri+1].cells[ci]
        c.text = ''
        r = c.paragraphs[0].add_run(val)
        r.font.size = Pt(8.5)
        if ci == 0:
            r.bold = True
            set_cell_bg(c, "EBF5FB")

doc.add_page_break()

# 양식 매핑표
add_styled_paragraph(doc, '7.3 품목-양식코드 자동 매핑표', font_size=12, bold=True, color=C_PRIMARY, space_before=8)

form_map = doc.add_table(rows=14, cols=5)
form_map.style = 'Table Grid'
form_map.alignment = WD_TABLE_ALIGNMENT.CENTER

fm_headers = ['품목코드', '품목명', '양식코드', '양식명', '검사항목 수']
for i, h in enumerate(fm_headers):
    c = form_map.rows[0].cells[i]
    c.text = ''
    r = c.paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = C_WHITE
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(c, "1B3A5C")

fm_data = [
    ('SM-GW-24', '글라스울(24K)', 'D122-1', '그라스울 인수검사', '8'),
    ('SM-CW-96', '세라믹차열재(96K)', 'D124-1', '세라믹울 96K', '7'),
    ('SM-CW-128', '세라믹차열재(128K)', 'D124-3', '세라믹울 120K+', '7'),
    ('SM-GI-I', '아연도금강판(I형)', 'D121-2', '방화소켓(벽체)', '9'),
    ('SM-GI-Z', '아연도금강판(Z형)', 'D121-2', '방화소켓(벽체)', '9'),
    ('SM-BRK-TB', '소켓 브라켓(상/하)', 'D121-2', '방화소켓(벽체)', '9'),
    ('SM-BRK-MD', '소켓 브라켓(중앙)', 'D121-2', '방화소켓(벽체)', '9'),
    ('SM-GP', '고정자재', 'D121-4', '방화플래싱', '7'),
    ('SM-SIL', '실란트', 'D104-1', 'EVA-EP100 양식', '4'),
    ('RM-MB', '난연컴파운드', 'D101-1', '배합원료 검사', '7'),
    ('RM-EA', 'EVA-EA33045', 'D103-1', '배합원료 검사', '4'),
    ('RM-EG50', '팽창흑연 #50', 'D102-1', '배합원료 검사', '7'),
    ('RM-EP', 'EVA-EP100', 'D104-1', '배합원료 검사', '4'),
]

for ri, row_data in enumerate(fm_data):
    for ci, val in enumerate(row_data):
        c = form_map.rows[ri+1].cells[ci]
        c.text = ''
        r = c.paragraphs[0].add_run(val)
        r.font.size = Pt(8)
        if ci == 0:
            r.bold = True
            set_cell_bg(c, "EBF5FB")

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 8. STEP 5 - 인수검사
# ═══════════════════════════════════════════════════════

doc.add_heading('8. [STEP 5] 인수검사 실행', level=1)

add_styled_paragraph(doc, '8.1 업무 개요', font_size=12, bold=True, color=C_PRIMARY, space_before=6)
add_styled_paragraph(doc,
    '입고 등록 시 자동 생성된 인수검사에 대해, 품질팀 실무자가 실제 측정을 수행하고 '
    '측정값을 시스템에 입력한다. 시스템은 입력된 측정값을 사규 기준치(cert_standard)와 '
    '자동 비교하여 항목별 PASS/FAIL을 판정하고, 전체 종합 판정을 산출한다.',
    font_size=10, space_after=6)

make_flow_table(doc, [
    {'no': '5-1', 'name': '검사 대상 확인', 'actor': '품질팀 실무자',
     'action': 'MES [인수검사] 메뉴에서\nPENDING 상태 검사 목록 확인\n- LOT번호, 품목, 양식 확인',
     'system': 'GET /api/inspections\nresult = "PENDING" 필터',
     'output': '검사 대상 목록', 'type': 'start'},
    {'no': '5-2', 'name': '시료 채취', 'actor': '품질팀 실무자',
     'action': '입고 자재에서 랜덤 3개 시료 채취\n(n=3, c=0 원칙)\n- 동일 LOT 내 서로 다른 위치에서 추출',
     'system': '-',
     'output': '시료 3개\n(n1, n2, n3)', 'type': 'process'},
    {'no': '5-3', 'name': '측정 실행', 'actor': '품질팀 실무자',
     'action': '검사항목별 측정 수행:\n- 줄자: 길이/너비/높이 (mm 정수)\n- 버니어캘리퍼스: 두께 (소수2자리)\n- 마이크로미터: 두께 (소수2자리)\n- 육안: 겉모양 확인 (OK/NG)\n- 저울: 질량 측정 (밀도 계산용)',
     'system': '-',
     'output': '측정 원시 데이터', 'type': 'process'},
    {'no': '5-4', 'name': '측정값 시스템 입력', 'actor': '품질팀 실무자',
     'action': '검사 상세 화면에서 측정값 입력:\n- 각 항목별 n1, n2, n3 값 입력\n- 성적서 참조항목은 n1만 입력\n- 비고란: 공인기관명, 발행일 기재',
     'system': 'PATCH /inspections/:id/details\nmeasured_n1, n2, n3 저장',
     'output': '측정 데이터 입력', 'type': 'process'},
    {'no': '5-5', 'name': '자동 판정', 'actor': '시스템 (자동)',
     'action': '시스템 자동 판정 로직:\n\u2460 항목별: cert_standard 대비 n1,n2,n3 비교\n  - 줄자/캘리퍼스: MIN(n1,n2,n3) >= 기준\n  - 육안: 1 = PASS\n  - 성적서: n1 존재 시 PASS\n\u2461 종합: 전 항목 PASS면 PASS, 1개라도 FAIL이면 FAIL',
     'system': 'judgeDetailItem()\ncalculateOverallResult()\ninspection.result 갱신',
     'output': 'PASS 또는 FAIL', 'type': 'decision'},
    {'no': '5-6', 'name': '결과 확인/서명', 'actor': '관리자 (팀장)',
     'action': '검사 결과 최종 확인:\n- 측정값 타당성 검토\n- n1=n2=n3 동일값 없는지 확인\n- 판정 결과 승인',
     'system': '검사결과 확정\n성적서 출력 가능',
     'output': '검사 완료\n\u2192 생산 투입 가능', 'type': 'end'},
], title='인수검사 실행 상세 플로우')

doc.add_page_break()

# 그라스울 검사 예시
add_styled_paragraph(doc, '8.2 검사 양식별 상세 예시: 그라스울(D122-1)', font_size=12, bold=True,
                     color=C_PRIMARY, space_before=6)

gw_table = doc.add_table(rows=9, cols=7)
gw_table.style = 'Table Grid'
gw_table.alignment = WD_TABLE_ALIGNMENT.CENTER

gw_headers = ['No', '검사항목', '기준', '검사방법', 'n1', 'n2', 'n3']
for i, h in enumerate(gw_headers):
    c = gw_table.rows[0].cells[i]
    c.text = ''
    r = c.paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = C_WHITE
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(c, "1B3A5C")

gw_data = [
    ('1', '겉모양', '한도견본 기준', '육안', 'OK', 'OK', 'OK'),
    ('2', '밀도', '\u226524 kg/m\u00b3', '질량\u00f7부피', '25.3', '24.8', '25.1'),
    ('3', '두께', '\u226525 mm', '버니어캘리퍼스', '26', '25', '26'),
    ('4', '너비', '\u22651400 mm', '줄자', '1402', '1401', '1403'),
    ('5', '열전도율(20\u2103)', '\u22640.037', '성적서', '0.035', '-', '-'),
    ('6', '열전도율(70\u2103)', '\u22640.048', '성적서', '0.046', '-', '-'),
    ('7', '열간수축온도', '\u2265300\u2103', '성적서', '350', '-', '-'),
    ('8', '공인시험', '1회/년', '공인기관', '\u2713', '-', '-'),
]

for ri, row_data in enumerate(gw_data):
    for ci, val in enumerate(row_data):
        c = gw_table.rows[ri+1].cells[ci]
        c.text = ''
        r = c.paragraphs[0].add_run(val)
        r.font.size = Pt(8.5)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if ci >= 4 and val not in ('-', ''):
            r.bold = True
            if val in ('OK', '\u2713'):
                r.font.color.rgb = C_SUCCESS

doc.add_paragraph()

make_info_box(doc, '\u26A0  인수검사 핵심 원칙 (C-701 Rev.5)', [
    '\u2022 n1, n2, n3는 동일 LOT 내 랜덤 3개 시료의 독립 측정값',
    '\u2022 세 값이 모두 동일하면 부적합 의심 대상 (NC-B: 측정 신뢰성 의심)',
    '\u2022 줄자 측정: mm 정수 기재 / 버니어캘리퍼스,마이크로미터: 소수 2자리',
    '\u2022 성적서 대체 항목: n1에 시험값 또는 성적서번호 기재, n2/n3는 빈칸 또는 N/A',
    '\u2022 판정: 합격 또는 불합격 중 하나만 체크 (양쪽 동시 체크 불가)',
])

doc.add_page_break()

# 판정 로직 상세
add_styled_paragraph(doc, '8.3 자동 판정 로직 상세', font_size=12, bold=True, color=C_PRIMARY, space_before=6)

judge_table = doc.add_table(rows=5, cols=4)
judge_table.style = 'Table Grid'
judge_table.alignment = WD_TABLE_ALIGNMENT.CENTER

jt_headers = ['검사방법', '판정 조건', '결과', '비고']
for i, h in enumerate(jt_headers):
    c = judge_table.rows[0].cells[i]
    c.text = ''
    r = c.paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = C_WHITE
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(c, "1B3A5C")

jt_data = [
    ('줄자/캘리퍼스/\n마이크로미터', 'MIN(n1,n2,n3) >= cert_standard', 'PASS', '3개 중 최솟값이\n기준 이상'),
    ('육안', 'n1=1 (OK표기 = 1)', 'PASS', 'OK=1, NG=0'),
    ('성적서', 'n1 값이 존재 (not null)', 'PASS', '성적서 번호 또는\n시험값 입력'),
    ('공인기관', 'n1 값이 존재', 'PASS', '1회/년 유효\n성적서 확인'),
]
for ri, (m, cond, result, note) in enumerate(jt_data):
    for ci, val in enumerate([m, cond, result, note]):
        c = judge_table.rows[ri+1].cells[ci]
        c.text = ''
        r = c.paragraphs[0].add_run(val)
        r.font.size = Pt(8.5)
        if ci == 2:
            r.bold = True
            r.font.color.rgb = C_SUCCESS
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

add_styled_paragraph(doc, '종합 판정 규칙:', font_size=10, bold=True, space_before=6)
rules = [
    '전 항목이 PASS이면 \u2192 종합 PASS',
    '1개 이상 FAIL이면 \u2192 종합 FAIL',
    '일부 항목이 아직 NA(미입력)이면 \u2192 종합 PENDING (검사 진행 중)',
]
for rule in rules:
    p = doc.add_paragraph(rule, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 9. STEP 6 - 작업지시
# ═══════════════════════════════════════════════════════

doc.add_heading('9. [STEP 6] 작업지시 및 생산', level=1)

add_styled_paragraph(doc, '9.1 업무 개요', font_size=12, bold=True, color=C_PRIMARY, space_before=6)
add_styled_paragraph(doc,
    '인수검사 완료된 자재를 투입하여 생산을 진행한다. '
    '작업지시는 BOM 구조에 따라 4개 공정(배합\u2192압출\u2192재단\u2192조립)이 자동 생성되며, '
    '각 공정의 작업자, 시간, 생산량 등의 실적을 입력한다.',
    font_size=10, space_after=6)

# 공정 플로우
add_styled_paragraph(doc, '9.2 생산 공정 플로우 (4단계 CASCADE)', font_size=12, bold=True, color=C_PRIMARY, space_before=8)

make_arrow_flow(doc, ['MIX\n(배합)', 'EXT\n(압출)', 'CUT\n(재단)', 'ASM\n(조립)'], "27AE60")

doc.add_paragraph()

wo_detail = doc.add_table(rows=5, cols=5)
wo_detail.style = 'Table Grid'
wo_detail.alignment = WD_TABLE_ALIGNMENT.CENTER

wo_headers = ['공정', '작업내용', '투입 자재', '산출물', '작업지시 수']
for i, h in enumerate(wo_headers):
    c = wo_detail.rows[0].cells[i]
    c.text = ''
    r = c.paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = C_WHITE
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(c, "1B3A5C")

wo_data = [
    ('MIX (배합)', '인정배합비율에 따라\n4종 원자재 혼합', 'RM-MB, RM-EA,\nRM-EP, RM-EG50', 'SA-MIX-MB\n(배합물 446kg)', '1건'),
    ('EXT (압출)', '배합물을 압출기로\n차열시트 성형', 'SA-MIX-MB', 'SA-EXT-5190\nSA-EXT-5125\nSA-EXT-4125', '3건'),
    ('CUT (재단)', '압출 시트를\n규격별로 재단', 'SA-EXT-xxxx', 'SA-CUT-SK (324매)\nSA-CUT-FL (167매)', '2건'),
    ('ASM (조립)', '소켓/플래싱/시트\n최종 조립', '재단 차열시트\n+ 부자재(강판,브라켓\n세라믹,글라스울 등)', 'FP-VA064 (11개)\nFP-FL-I (99장)\n등 완제품 7종', '7건'),
]
for ri, row_data in enumerate(wo_data):
    for ci, val in enumerate(row_data):
        c = wo_detail.rows[ri+1].cells[ci]
        c.text = ''
        r = c.paragraphs[0].add_run(val)
        r.font.size = Pt(8.5)
        if ci == 0:
            r.bold = True
            set_cell_bg(c, "D5F5E3")

doc.add_paragraph()

make_flow_table(doc, [
    {'no': '6-1', 'name': '작업지시 자동 생성', 'actor': '생산팀 (또는 자재팀)',
     'action': '수주 상세 \u2192 [작업지시 생성] 클릭\n- BOM 기반 4공정 13건 자동 생성\n- 작업일자: 납기일 -7일 자동 설정',
     'system': 'POST /orders/:id/\ngenerate-work-orders\nWO-{공정}-YYMMDD-NNN',
     'output': '작업지시 13건\n(PLANNED)', 'type': 'process'},
    {'no': '6-2', 'name': '작업 배정', 'actor': '생산팀 관리자',
     'action': '작업지시별 작업자 배정:\n- 오전/오후/야간 근무자 지정\n- 설비(라인) 배정',
     'system': 'PATCH /work-orders/:id\nam_worker, pm_worker 등',
     'output': '작업 배정 완료', 'type': 'process'},
    {'no': '6-3', 'name': '생산 실행', 'actor': '생산팀 작업자',
     'action': '작업지시서에 따라 생산 실행:\n- 원자재 투입량 기록\n- 생산 시작/종료 시간 기록\n- 불량/로스 기록',
     'system': 'PLANNED \u2192 IN_PROGRESS',
     'output': '생산 실적 데이터', 'type': 'process'},
    {'no': '6-4', 'name': '실적 입력/완료', 'actor': '생산팀 작업자',
     'action': '실제 생산량, 투입량, 로스 입력\n- 중간검사 결과 기록\n- LOT번호 연결',
     'system': 'IN_PROGRESS \u2192 COMPLETED\nactual_qty, scrap_kg 등',
     'output': '작업 완료\n\u2192 출하 준비', 'type': 'end'},
], title='작업지시 실행 상세 플로우')

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 10. 실제 운영 예시
# ═══════════════════════════════════════════════════════

doc.add_heading('10. 실제 운영 예시 (SO-260303-001)', level=1)

add_styled_paragraph(doc,
    '아래는 실제 시스템에서 처리된 수주 SO-260303-001의 전체 프로세스 진행 현황이다.',
    font_size=10, space_after=8)

# 수주 정보
add_styled_paragraph(doc, '10.1 수주 정보', font_size=12, bold=True, color=C_PRIMARY, space_before=6)

order_info = doc.add_table(rows=5, cols=2)
order_info.style = 'Table Grid'
order_info.alignment = WD_TABLE_ALIGNMENT.CENTER
order_data = [
    ('수주번호', 'SO-260303-001'),
    ('고객', '㈜하나로엔지니어링'),
    ('현장', '인천검단101역세권C1현장'),
    ('총 수량', '22세트 (10개 라인, 4개 구조)'),
    ('현재 상태', 'IN_PRODUCTION (생산진행 중)'),
]
for i, (k, v) in enumerate(order_data):
    c0 = order_info.rows[i].cells[0]
    c0.text = ''
    r0 = c0.paragraphs[0].add_run(k)
    r0.bold = True
    r0.font.size = Pt(9)
    set_cell_bg(c0, "EBF5FB")
    c0.width = Cm(3)
    c1 = order_info.rows[i].cells[1]
    c1.text = ''
    r1 = c1.paragraphs[0].add_run(v)
    r1.font.size = Pt(9)

doc.add_paragraph()

# 구조별 내역
add_styled_paragraph(doc, '10.2 수주 구조별 내역', font_size=12, bold=True, color=C_PRIMARY, space_before=6)

struct_table = doc.add_table(rows=6, cols=5)
struct_table.style = 'Table Grid'
struct_table.alignment = WD_TABLE_ALIGNMENT.CENTER

sh = ['구조코드', '유형', '관통부 치수', '소켓수(N)', '세트수']
for i, h in enumerate(sh):
    c = struct_table.rows[0].cells[i]
    c.text = ''
    r = c.paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = C_WHITE
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(c, "1B3A5C")

sd = [
    ('VA-064', '벽체', '550x250 ~ 1700x550', '1', '11'),
    ('VT-01', '벽체', '300x200', '2', '3'),
    ('HTG-1.69', '입상', '1550x700 ~ 1200x900', '2', '3'),
    ('HTG-064', '입상', '700x700 ~ 1200x850', '2', '5'),
    ('합계', '', '', '', '22'),
]
for ri, row_data in enumerate(sd):
    for ci, val in enumerate(row_data):
        c = struct_table.rows[ri+1].cells[ci]
        c.text = ''
        r = c.paragraphs[0].add_run(val)
        r.font.size = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if ri == 4:
            r.bold = True
            set_cell_bg(c, "FDEBD0")

doc.add_paragraph()

# 발주/입고/검사 현황
add_styled_paragraph(doc, '10.3 발주/입고/인수검사 현황', font_size=12, bold=True, color=C_PRIMARY, space_before=6)

pr_table = doc.add_table(rows=14, cols=6)
pr_table.style = 'Table Grid'
pr_table.alignment = WD_TABLE_ALIGNMENT.CENTER

pr_h = ['품목코드', '품목명', '발주수량', '입고상태', '자체LOT', '검사결과']
for i, h in enumerate(pr_h):
    c = pr_table.rows[0].cells[i]
    c.text = ''
    r = c.paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = C_WHITE
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(c, "1B3A5C")

pr_data = [
    ('RM-EA', 'EVA-EA33045', '6.94 kg', 'RECEIVED', 'EZ-260330-009', 'PENDING'),
    ('RM-EG50', '팽창흑연 #50', '9.25 kg', 'RECEIVED', 'EZ-260330-010', 'PENDING'),
    ('RM-EP', 'EVA-EP100', '6.94 kg', 'RECEIVED', 'EZ-260330-011', 'PENDING'),
    ('RM-MB', '난연컴파운드', '23.13 kg', 'RECEIVED', 'EZ-260330-012', 'PENDING'),
    ('SM-BRK-MD', '브라켓(중앙)', '24 EA', 'RECEIVED', 'EZ-260330-013', 'PASS'),
    ('SM-BRK-TB', '브라켓(상/하)', '112 EA', 'RECEIVED', 'EZ-260330-008', 'PASS'),
    ('SM-CW-128', '세라믹울(128K)', '37.18 M', 'RECEIVED', 'EZ-260330-014', 'PENDING'),
    ('SM-CW-96', '세라믹울(96K)', '198.59 M', 'RECEIVED', 'EZ-260330-006', 'PASS'),
    ('SM-GI-I', '강판(I형)', '99 EA', 'RECEIVED', 'EZ-260330-007', 'PASS'),
    ('SM-GI-Z', '강판(Z형)', '68 EA', 'RECEIVED', 'EZ-260330-015', 'PENDING'),
    ('SM-GP', '고정자재', '60 EA', 'RECEIVED', 'EZ-260330-016', 'PENDING'),
    ('SM-GW-24', '글라스울(24K)', '225.73 M', 'RECEIVED', 'EZ-260330-005', 'PASS'),
    ('SM-SIL', '실란트', '20 EA', 'RECEIVED', 'EZ-260330-017', 'PENDING'),
]

for ri, row_data in enumerate(pr_data):
    for ci, val in enumerate(row_data):
        c = pr_table.rows[ri+1].cells[ci]
        c.text = ''
        r = c.paragraphs[0].add_run(val)
        r.font.size = Pt(8)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if ci == 5:
            if val == 'PASS':
                r.bold = True
                r.font.color.rgb = C_SUCCESS
                set_cell_bg(c, "D5F5E3")
            elif val == 'PENDING':
                r.font.color.rgb = C_WARNING

doc.add_page_break()

# 작업지시 현황
add_styled_paragraph(doc, '10.4 작업지시 현황 (13건)', font_size=12, bold=True, color=C_PRIMARY, space_before=6)

wo_table = doc.add_table(rows=14, cols=5)
wo_table.style = 'Table Grid'
wo_table.alignment = WD_TABLE_ALIGNMENT.CENTER

wo_h = ['작업지시번호', '공정', '계획수량', '내용', '상태']
for i, h in enumerate(wo_h):
    c = wo_table.rows[0].cells[i]
    c.text = ''
    r = c.paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = C_WHITE
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(c, "1B3A5C")

wo_rows = [
    ('WO-MIX-20260406-001', 'MIX', '446.26 kg', '4종 원자재 배합', 'PLANNED'),
    ('WO-EXT-20260406-001', 'EXT', '192.05 m', '플래싱 차열시트 4T(Z형)', 'PLANNED'),
    ('WO-EXT-20260406-002', 'EXT', '192.05 m', '플래싱 차열시트 5T(I형)', 'PLANNED'),
    ('WO-EXT-20260406-003', 'EXT', '372.60 m', '소켓용 차열시트 5T-190', 'PLANNED'),
    ('WO-CUT-20260406-001', 'CUT', '167 매', '플래싱용 재단', 'PLANNED'),
    ('WO-CUT-20260406-002', 'CUT', '324 매', '소켓용 재단', 'PLANNED'),
    ('WO-ASM-20260406-001', 'ASM', '99 EA', '플래싱(I형) 조립', 'PLANNED'),
    ('WO-ASM-20260406-002', 'ASM', '68 EA', '플래싱(Z형) 조립', 'PLANNED'),
    ('WO-ASM-20260406-003', 'ASM', '60 EA', '틈새복합시트 조립', 'PLANNED'),
    ('WO-ASM-20260406-004', 'ASM', '5 EA', '방화소켓(HTG-064)', 'PLANNED'),
    ('WO-ASM-20260406-005', 'ASM', '6 EA', '방화소켓(HTG-1.69)', 'PLANNED'),
    ('WO-ASM-20260406-006', 'ASM', '11 EA', '방화소켓(VA-064)', 'PLANNED'),
    ('WO-ASM-20260406-007', 'ASM', '6 EA', '방화소켓(VT-01)', 'PLANNED'),
]
for ri, row_data in enumerate(wo_rows):
    for ci, val in enumerate(row_data):
        c = wo_table.rows[ri+1].cells[ci]
        c.text = ''
        r = c.paragraphs[0].add_run(val)
        r.font.size = Pt(8)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if ci == 1:
            r.bold = True
            colors = {'MIX': 'E8DAEF', 'EXT': 'D4EFDF', 'CUT': 'FDEBD0', 'ASM': 'D6EAF8'}
            set_cell_bg(c, colors.get(val, 'FFFFFF'))

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 11. 부적합 처리 플로우
# ═══════════════════════════════════════════════════════

doc.add_heading('11. 부적합 처리 플로우', level=1)

add_styled_paragraph(doc,
    '인수검사 결과 FAIL 판정 시, 아래 절차에 따라 부적합 처리를 진행한다.',
    font_size=10, space_after=8)

make_vertical_flow(doc, [
    {'label': '인수검사 FAIL 판정', 'detail': '1개 이상 항목 기준 미달', 'type': 'start', 'color': 'FADBD8'},
    {'label': '부적합 코드 분류', 'detail': 'NC-A(기준미달) / NC-B(신뢰성의심) / NC-C(기재누락) / NC-D(양식불일치)', 'type': 'process', 'color': 'FDEBD0'},
    {'label': '관리자 판단', 'detail': '재검사 / 반품 / 특채(조건부 합격) 결정', 'type': 'decision', 'color': 'FDEBD0'},
    {'label': '시정조치 실행', 'detail': 'NC-A: 반품 또는 재검사 | NC-B: 재측정 후 성적서 재작성 | NC-C: 보완기재 | NC-D: 양식교체', 'type': 'process', 'color': 'D6EAF8'},
    {'label': '조치 완료 / 기록 보존', 'detail': '부적합 보고서 작성, 시정조치 기록', 'type': 'end', 'color': 'D5F5E3'},
])

doc.add_paragraph()

# 부적합 코드
nc_table = doc.add_table(rows=5, cols=4)
nc_table.style = 'Table Grid'
nc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
nc_h = ['코드', '유형', '설명', '조치']
for i, h in enumerate(nc_h):
    c = nc_table.rows[0].cells[i]
    c.text = ''
    r = c.paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = C_WHITE
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(c, "E74C3C")

nc_data = [
    ('NC-A', '기준치 미달', '측정값이 사규 기준 미만', '재검사 또는 반품 처리'),
    ('NC-B', '측정 신뢰성 의심', 'n1=n2=n3 동일값, 단위오류', '재측정 후 성적서 재작성'),
    ('NC-C', '기재 누락/오류', '날짜, LOT, 검사자 등 누락', '보완 기재'),
    ('NC-D', '양식/성적서 불일치', '잘못된 양식, 성적서 만료', '양식 교체 또는 성적서 재발급'),
]
for ri, row_data in enumerate(nc_data):
    for ci, val in enumerate(row_data):
        c = nc_table.rows[ri+1].cells[ci]
        c.text = ''
        r = c.paragraphs[0].add_run(val)
        r.font.size = Pt(8.5)
        if ci == 0:
            r.bold = True
            r.font.color.rgb = C_DANGER

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 12. 부록
# ═══════════════════════════════════════════════════════

doc.add_heading('12. 부록: 시스템 API 엔드포인트 요약', level=1)

add_styled_paragraph(doc,
    'MES 시스템의 주요 API 엔드포인트와 업무 매핑을 정리한다.',
    font_size=10, space_after=8)

api_table = doc.add_table(rows=11, cols=4)
api_table.style = 'Table Grid'
api_table.alignment = WD_TABLE_ALIGNMENT.CENTER

api_h = ['업무단계', 'HTTP', 'API 경로', '설명']
for i, h in enumerate(api_h):
    c = api_table.rows[0].cells[i]
    c.text = ''
    r = c.paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = C_WHITE
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(c, "1B3A5C")

api_data = [
    ('수주등록', 'POST', '/api/orders', '수주 신규 생성'),
    ('BOM전개', 'POST', '/api/orders/:id/explode-bom', '치수 기반 BOM 소요량 계산'),
    ('발주서 생성', 'POST', '/api/orders/:id/create-pr', '부족 자재 발주서 자동 생성'),
    ('발주 상태변경', 'PATCH', '/api/purchase-requests/:id/status', 'DRAFT\u2192SUBMITTED\u2192APPROVED\u2192ORDERED'),
    ('입고 등록', 'POST', '/api/purchase-requests/:prId/items/:priId/receive', 'LOT 생성 + 인수검사 자동 연결'),
    ('검사 측정입력', 'PATCH', '/api/inspections/:id/details', '측정값 입력 + 자동 판정'),
    ('작업지시 생성', 'POST', '/api/orders/:id/generate-work-orders', '4공정 CASCADE 자동 생성'),
    ('입고 현황', 'GET', '/api/purchase-requests/:prId/receiving-status', '발주서별 입고/검사 현황'),
    ('검사 목록', 'GET', '/api/inspections', '인수검사 목록 조회'),
    ('작업지시 목록', 'GET', '/api/work-orders', '작업지시 목록 조회'),
]

for ri, row_data in enumerate(api_data):
    for ci, val in enumerate(row_data):
        c = api_table.rows[ri+1].cells[ci]
        c.text = ''
        r = c.paragraphs[0].add_run(val)
        r.font.size = Pt(8)
        if ci == 1:
            r.bold = True
            colors = {'POST': 'D5F5E3', 'PATCH': 'FDEBD0', 'GET': 'D6EAF8'}
            set_cell_bg(c, colors.get(val, 'FFFFFF'))
        if ci == 2:
            r.font.name = 'Consolas'
            r.font.size = Pt(7.5)

doc.add_paragraph()
doc.add_paragraph()

# 마지막 페이지 - 개정이력
add_styled_paragraph(doc, '개정 이력', font_size=14, bold=True, color=C_PRIMARY,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=20)

rev_table = doc.add_table(rows=3, cols=4)
rev_table.style = 'Table Grid'
rev_table.alignment = WD_TABLE_ALIGNMENT.CENTER

rev_h = ['Rev', '일자', '내용', '작성자']
for i, h in enumerate(rev_h):
    c = rev_table.rows[0].cells[i]
    c.text = ''
    r = c.paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = C_WHITE
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(c, "1B3A5C")

rev_data = [
    ('0', '2026-03-30', '최초 작성', 'June'),
    ('1', '2026-03-30', '브라켓/강판 BOM 추가 반영, 인수검사 양식매핑 보완', 'June'),
]
for ri, row_data in enumerate(rev_data):
    for ci, val in enumerate(row_data):
        c = rev_table.rows[ri+1].cells[ci]
        c.text = ''
        r = c.paragraphs[0].add_run(val)
        r.font.size = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ─── 끝 표시 ───
doc.add_paragraph()
p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_end = p_end.add_run('--- 끝 ---')
r_end.font.size = Pt(12)
r_end.font.color.rgb = C_GRAY
r_end.italic = True


# ═══════════════════════════════════════════════════════
# 저장
# ═══════════════════════════════════════════════════════

output_path = os.path.join(os.path.dirname(__file__), 'GDL_20260330_업무프로세스플로우가이드.docx')
doc.save(output_path)
print(f'문서 생성 완료: {output_path}')
print(f'파일 크기: {os.path.getsize(output_path) / 1024:.1f} KB')
